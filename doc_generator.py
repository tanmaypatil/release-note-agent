import os
from datetime import datetime
from docx import Document
from docx.shared import Pt
import db

DOCS_DIR = os.path.dirname(os.path.abspath(__file__))


def _doc_path(label):
    return os.path.join(DOCS_DIR, f"release_notes_{label}.docx")


def create_doc(label):
    defects = db.get_resolved_defects(label)
    release_info = db.get_release_info(label)

    doc = Document()

    # Title heading
    doc.add_heading(f'Release Notes — {label}', level=1)

    # Metadata line
    build_no = release_info['build_no'] if release_info else 'N/A'
    generated_date = datetime.now().strftime('%Y-%m-%d')
    doc.add_paragraph(f'Build: {build_no}  |  Generated: {generated_date}')

    # 4-column table
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'

    # Header row (bold)
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(['ID', 'Title', 'Date', 'Developer Comment']):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].bold = True

    # Data rows
    for defect in defects:
        row_cells = table.add_row().cells
        row_cells[0].text = str(defect['id'])
        row_cells[1].text = defect['title']
        row_cells[2].text = defect['date']
        row_cells[3].text = defect.get('developer_comment', '')

    path = _doc_path(label)
    doc.save(path)
    return path


def remove_row(label, defect_id):
    path = _doc_path(label)
    if not os.path.exists(path):
        return

    doc = Document(path)
    for table in doc.tables:
        for row in table.rows:
            if row.cells[0].text == str(defect_id):
                table._tbl.remove(row._tr)
                break

    doc.save(path)
