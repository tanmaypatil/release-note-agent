"""
Section 2 — Doc generator tests.

Requires test data already in the DB (run test_db.py first, or the script
seeds its own data if the DB is empty).

Run from the project root:
    python tests/test_doc.py
"""
import sys
import os
import tempfile

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

import db
import doc_generator
from docx import Document

# Use a temporary database so each run starts from a clean slate
_tmp = tempfile.NamedTemporaryFile(suffix='.db', delete=False)
db.DB_PATH = _tmp.name
_tmp.close()

LABEL = 'v1.0'


def ensure_seed_data():
    db.init_db()
    resolved = db.get_resolved_defects(LABEL)
    if len(resolved) < 2:
        db.create_defect('Login crashes on Safari', '2024-01-15', 'Fixed null pointer', LABEL, 'RESOLVED')
        id2 = db.create_defect('Tooltip not showing', '2024-01-16', 'Updated z-index', LABEL, 'RESOLVED')
        db.create_defect('Password reset delayed', '2024-01-17', '', LABEL, 'OPEN')
        resolved = db.get_resolved_defects(LABEL)
    return resolved


def test_create_doc(resolved):
    path = doc_generator.create_doc(LABEL)
    assert os.path.exists(path), f'Doc not found at {path}'

    doc = Document(path)
    rows = list(doc.tables[0].rows)
    expected = len(resolved) + 1  # data rows + header
    assert len(rows) == expected, f'Expected {expected} rows, got {len(rows)}'
    print(f'create_doc:          PASS  ({len(rows)-1} data rows + header)')
    return path, rows


def test_remove_row(resolved):
    remove_id = resolved[0]['id']
    doc_generator.remove_row(LABEL, remove_id)

    doc = Document(doc_generator._doc_path(LABEL))
    rows = list(doc.tables[0].rows)
    ids_in_doc = [row.cells[0].text for row in rows[1:]]
    assert str(remove_id) not in ids_in_doc, f'Defect {remove_id} should have been removed'
    print(f'remove_row:          PASS  (row {remove_id} gone, {len(rows)-1} data rows remain)')


def test_remove_row_noop():
    doc_generator.remove_row('does-not-exist', 999)
    print('remove_row (no-op):  PASS')


def teardown():
    path = doc_generator._doc_path(LABEL)
    if os.path.exists(path):
        os.remove(path)
    os.unlink(db.DB_PATH)


def main():
    print('=== Doc Generator ===')
    resolved = ensure_seed_data()
    _, rows = test_create_doc(resolved)
    test_remove_row(resolved)
    test_remove_row_noop()
    teardown()
    print('\nAll doc generator tests passed.')


if __name__ == '__main__':
    main()
